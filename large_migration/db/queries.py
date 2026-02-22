"""
Query tools for the migration multi-agent system.
All tools are decorated with @function_tool for use with OpenAI Agents SDK.
"""

import json
import os
import re
from pathlib import Path
from typing import Optional, Tuple
from agents import function_tool
from .connection import get_connection, init_database

# Import logger
try:
    from utils.logger import get_logger
    logger = get_logger()
except ImportError:
    logger = None

def _log(level: str, message: str, details: str = None):
    """Helper to log if logger is available."""
    if logger:
        method = getattr(logger, level, logger.info)
        method(message, details)

# Input folder path for user-provided files
INPUT_DIR = Path(__file__).parent.parent / "input"


def _close_cursor(cursor, conn):
    """Helper to safely close cursor and connection."""
    if cursor:
        cursor.close()
    if conn:
        conn.close()


def _parse_file_content(file_path: str) -> Tuple[str, str]:
    """
    Parse file content based on file extension.
    Supports: TXT, PDF, DOCX, CSV, XLSX, XLS, JSON, YAML, SQL, MD

    Args:
        file_path: Path to the file

    Returns:
        Tuple of (content, file_type)
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    try:
        # Plain text files
        if ext in ['.txt', '.sql', '.md', '.json', '.yaml', '.yml']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read(), ext[1:]  # Remove leading dot

        # PDF files
        elif ext == '.pdf':
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {i+1} ---\n{page_text}")
            return "\n\n".join(text_parts), "pdf"

        # Word documents
        elif ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # Also extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    paragraphs.append("\n[Table]\n" + "\n".join(table_text))
            return "\n\n".join(paragraphs), "docx"

        # Excel files (XLSX)
        elif ext == '.xlsx':
            import pandas as pd
            # Read all sheets
            xlsx = pd.ExcelFile(file_path)
            all_sheets = []
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                sheet_content = f"--- Sheet: {sheet_name} ---\n"
                sheet_content += df.to_string(index=False)
                all_sheets.append(sheet_content)
            return "\n\n".join(all_sheets), "xlsx"

        # Legacy Excel files (XLS)
        elif ext == '.xls':
            import pandas as pd
            xlsx = pd.ExcelFile(file_path, engine='xlrd')
            all_sheets = []
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                sheet_content = f"--- Sheet: {sheet_name} ---\n"
                sheet_content += df.to_string(index=False)
                all_sheets.append(sheet_content)
            return "\n\n".join(all_sheets), "xls"

        # CSV files
        elif ext == '.csv':
            import pandas as pd
            df = pd.read_csv(file_path)
            return df.to_string(index=False), "csv"

        # Unknown format - try as text
        else:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read(), "text"
            except UnicodeDecodeError:
                return f"[Binary file - cannot parse {ext} format]", "binary"

    except ImportError as e:
        return f"[Error: Missing library to parse {ext} files: {e}]", "error"
    except Exception as e:
        return f"[Error parsing file: {str(e)}]", "error"


# =============================================================================
# INPUT FILE TOOLS (List and read files from input folder)
# =============================================================================

@function_tool
def list_input_files() -> str:
    """
    List all files available in the input folder.
    Users should place their input files (vendor templates, mapping documents,
    dbt models) in the input folder for processing.

    Supported formats: TXT, PDF, DOCX, CSV, XLSX, XLS, JSON, YAML, SQL, MD

    Returns:
        List of files in the input folder with their sizes and types.
    """
    _log("tool", "Listing input files", "Scanning input/ folder")
    try:
        if not INPUT_DIR.exists():
            INPUT_DIR.mkdir(parents=True, exist_ok=True)
            return "Input folder is empty. Please add files to the 'input' folder."

        # File type mapping
        type_map = {
            '.csv': 'spreadsheet (CSV)',
            '.xlsx': 'spreadsheet (Excel)',
            '.xls': 'spreadsheet (Excel legacy)',
            '.sql': 'SQL/dbt model',
            '.txt': 'text',
            '.md': 'markdown',
            '.json': 'JSON',
            '.yaml': 'YAML',
            '.yml': 'YAML',
            '.pdf': 'PDF document',
            '.docx': 'Word document',
            '.doc': 'Word document (legacy)',
        }

        # Patterns that suggest dbt/SQL content in .txt files
        dbt_patterns = ['dbt', 'sql', 'add_info', 'stg_', 'prep', 'final']

        files = []
        for f in INPUT_DIR.iterdir():
            if f.is_file() and not f.name.startswith('.'):
                size = f.stat().st_size
                ext = f.suffix.lower()
                file_type = type_map.get(ext, "unknown")

                # Check if .txt file might be a dbt model based on filename
                if ext == '.txt':
                    name_lower = f.name.lower()
                    if any(pattern in name_lower for pattern in dbt_patterns):
                        file_type = "SQL/dbt model (.txt)"

                files.append({
                    "name": f.name,
                    "path": str(f),
                    "size": size,
                    "type": file_type,
                    "ext": ext
                })

        if not files:
            return "Input folder is empty. Please add files to the 'input' folder for processing."

        result = f"Files in input folder ({len(files)} files):\n\n"
        for f in sorted(files, key=lambda x: x["name"]):
            size_kb = f["size"] / 1024
            result += f"- {f['name']} ({f['type']}, {size_kb:.1f} KB)\n"

        result += "\nSupported formats: TXT, PDF, DOCX, CSV, XLSX, XLS, JSON, YAML, SQL, MD"
        result += "\n\nUse read_input_file() to preview contents, or load_vendor_template()/load_mapping_document() to load into database."
        return result
    except Exception as e:
        return f"Error listing input files: {str(e)}"


@function_tool
def read_input_file(filename: str) -> str:
    """
    Read and parse the contents of a file from the input folder.
    Supports: TXT, PDF, DOCX, CSV, XLSX, XLS, JSON, YAML, SQL, MD

    Use this to preview file contents before loading them into the database.

    Args:
        filename: Name of the file in the input folder (not full path)

    Returns:
        Parsed contents of the file (first 8000 characters if large).
    """
    _log("file_op", f"Reading file: {filename}", "Parsing content...")
    try:
        file_path = INPUT_DIR / filename

        if not file_path.exists():
            available = [f.name for f in INPUT_DIR.iterdir() if f.is_file() and not f.name.startswith('.')]
            return f"File '{filename}' not found in input folder. Available files: {', '.join(available) if available else 'none'}"

        # Parse file content based on type
        content, file_type = _parse_file_content(str(file_path))

        if file_type == "error":
            return f"Error parsing file '{filename}': {content}"

        # Truncate if too large
        max_chars = 8000
        if len(content) > max_chars:
            return f"File: {filename}\nType: {file_type}\nSize: {len(content)} characters\n\nFirst {max_chars} characters:\n\n{content[:max_chars]}\n\n... [truncated, {len(content) - max_chars} more characters]"

        return f"File: {filename}\nType: {file_type}\nSize: {len(content)} characters\n\nContent:\n\n{content}"
    except Exception as e:
        return f"Error reading file: {str(e)}"


@function_tool
def get_input_file_path(filename: str) -> str:
    """
    Get the full path for a file in the input folder.
    Use this to get the path for loading functions.

    Args:
        filename: Name of the file in the input folder

    Returns:
        Full path to the file.
    """
    try:
        file_path = INPUT_DIR / filename

        if not file_path.exists():
            available = [f.name for f in INPUT_DIR.iterdir() if f.is_file()]
            return f"File '{filename}' not found. Available files: {', '.join(available) if available else 'none'}"

        return str(file_path)
    except Exception as e:
        return f"Error: {str(e)}"


# =============================================================================
# DOCUMENT LOADING TOOLS (Intake Agent)
# =============================================================================

@function_tool
def load_vendor_template(file_path: str) -> str:
    """
    Load and store a vendor template document containing field definitions,
    types, business rules, and agreement-type specific requirements.

    Supports: TXT, PDF, DOCX, CSV, XLSX, XLS, JSON, YAML, MD

    Args:
        file_path: Path to the vendor template file

    Returns:
        Success message with template summary or error message.
    """
    _log("file_op", f"Loading vendor template", file_path)
    try:
        init_database()
        _log("db", "Database initialized")

        # Check if file_path is just a filename (look in input folder)
        path = Path(file_path)
        if not path.is_absolute() and not path.exists():
            path = INPUT_DIR / file_path

        if not path.exists():
            return f"Error: File not found at {file_path}"

        # Parse file content based on type
        content, file_type = _parse_file_content(str(path))

        if file_type == "error":
            return f"Error parsing file: {content}"

        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO documents (doc_type, doc_name, content, updated_at)
            VALUES ('vendor_template', %s, %s, CURRENT_TIMESTAMP)
            ON CONFLICT (doc_type, doc_name)
            DO UPDATE SET content = EXCLUDED.content, updated_at = CURRENT_TIMESTAMP
        """, (str(path), content))

        conn.commit()
        _close_cursor(cursor, conn)

        line_count = len(content.splitlines())
        return f"Vendor template loaded successfully.\nFile: {path.name}\nFormat: {file_type}\nLines: {line_count}\nCharacters: {len(content)}"
    except FileNotFoundError:
        return f"Error: File not found at {file_path}"
    except Exception as e:
        return f"Error loading vendor template: {str(e)}"


@function_tool
def load_mapping_document(file_path: str) -> str:
    """
    Load and store a mapping document containing source-to-target field mappings,
    transformation rules, and join key specifications.

    Supports: TXT, PDF, DOCX, CSV, XLSX, XLS, JSON, YAML, MD

    Args:
        file_path: Path to the mapping document

    Returns:
        Success message with mapping summary or error message.
    """
    try:
        init_database()

        # Check if file_path is just a filename (look in input folder)
        path = Path(file_path)
        if not path.is_absolute() and not path.exists():
            path = INPUT_DIR / file_path

        if not path.exists():
            return f"Error: File not found at {file_path}"

        # Parse file content based on type
        content, file_type = _parse_file_content(str(path))

        if file_type == "error":
            return f"Error parsing file: {content}"

        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO documents (doc_type, doc_name, content, updated_at)
            VALUES ('mapping_document', %s, %s, CURRENT_TIMESTAMP)
            ON CONFLICT (doc_type, doc_name)
            DO UPDATE SET content = EXCLUDED.content, updated_at = CURRENT_TIMESTAMP
        """, (str(path), content))

        conn.commit()
        _close_cursor(cursor, conn)

        line_count = len(content.splitlines())
        return f"Mapping document loaded successfully.\nFile: {path.name}\nFormat: {file_type}\nLines: {line_count}\nCharacters: {len(content)}"
    except FileNotFoundError:
        return f"Error: File not found at {file_path}"
    except Exception as e:
        return f"Error loading mapping document: {str(e)}"


@function_tool
def load_dbt_model(file_path: str, model_name: str, layer: str) -> str:
    """
    Load and store a dbt model file for analysis and modification.
    Accepts both .sql and .txt files containing SQL/dbt code.

    Args:
        file_path: Path to the dbt model file (.sql or .txt with SQL content)
        model_name: Name of the dbt model (e.g., 'ams_ai_dbt_add_info')
        layer: Model layer - 'prep' for preparation layer or 'final' for staging/final layer

    Returns:
        Success message with model summary or error message.
    """
    _log("file_op", f"Loading dbt model: {model_name} ({layer})", file_path)
    try:
        init_database()
        _log("db", "Database initialized")

        if layer not in ('prep', 'final'):
            return "Error: layer must be 'prep' or 'final'"

        # Check if file_path is just a filename (look in input folder)
        path = Path(file_path)
        if not path.is_absolute() and not path.exists():
            path = INPUT_DIR / file_path

        if not path.exists():
            _log("error", f"File not found: {file_path}")
            return f"Error: File not found at {file_path}"

        _log("file_op", f"Reading file: {path.name}")
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()

        conn = get_connection()
        cursor = conn.cursor()

        doc_name = f"{model_name}_{layer}"
        cursor.execute("""
            INSERT INTO documents (doc_type, doc_name, content, updated_at)
            VALUES ('dbt_model', %s, %s, CURRENT_TIMESTAMP)
            ON CONFLICT (doc_type, doc_name)
            DO UPDATE SET content = EXCLUDED.content, updated_at = CURRENT_TIMESTAMP
        """, (doc_name, content))

        # Also store in dbt_modifications for tracking changes
        cursor.execute("""
            INSERT INTO dbt_modifications (model_name, layer, original_sql, status)
            VALUES (%s, %s, %s, 'loaded')
            ON CONFLICT (model_name, layer)
            DO UPDATE SET original_sql = EXCLUDED.original_sql, status = 'loaded'
        """, (model_name, layer, content))

        conn.commit()
        _close_cursor(cursor, conn)

        line_count = len(content.splitlines())
        file_ext = path.suffix
        _log("success", f"dbt model loaded: {model_name} ({layer})", f"{line_count} lines from {file_ext} file")
        return f"dbt model '{model_name}' ({layer} layer) loaded successfully!\n\nFile: {path.name}\nFormat: {file_ext}\nLines: {line_count}\n\nThe model is now ready for analysis and modification."
    except FileNotFoundError:
        _log("error", f"File not found: {file_path}")
        return f"Error: File not found at {file_path}"
    except Exception as e:
        return f"Error loading dbt model: {str(e)}"


@function_tool
def save_document_to_store(doc_type: str, doc_name: str, content: str, parsed_data: Optional[str] = None) -> str:
    """
    Save a document or parsed data to the document store.

    Args:
        doc_type: Type of document (e.g., 'vendor_template', 'mapping_document', 'dbt_model')
        doc_name: Unique name for the document
        content: Raw content of the document
        parsed_data: Optional JSON string of parsed/structured data

    Returns:
        Success message or error message.
    """
    try:
        init_database()
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO documents (doc_type, doc_name, content, parsed_data, updated_at)
            VALUES (%s, %s, %s, %s, CURRENT_TIMESTAMP)
            ON CONFLICT (doc_type, doc_name)
            DO UPDATE SET content = EXCLUDED.content, parsed_data = EXCLUDED.parsed_data, updated_at = CURRENT_TIMESTAMP
        """, (doc_type, doc_name, content, parsed_data))

        conn.commit()
        _close_cursor(cursor, conn)

        return f"Document '{doc_name}' of type '{doc_type}' saved successfully."
    except Exception as e:
        return f"Error saving document: {str(e)}"


@function_tool
def get_document(doc_type: str, doc_name: Optional[str] = None) -> str:
    """
    Retrieve a document from the store.

    Args:
        doc_type: Type of document to retrieve
        doc_name: Optional specific document name. If not provided, returns latest of type.

    Returns:
        Document content and metadata or error message.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        if doc_name:
            cursor.execute("""
                SELECT doc_name, content, parsed_data, created_at, updated_at
                FROM documents WHERE doc_type = %s AND doc_name = %s
            """, (doc_type, doc_name))
        else:
            cursor.execute("""
                SELECT doc_name, content, parsed_data, created_at, updated_at
                FROM documents WHERE doc_type = %s ORDER BY updated_at DESC LIMIT 1
            """, (doc_type,))

        row = cursor.fetchone()
        _close_cursor(cursor, conn)

        if not row:
            return f"No document found for type '{doc_type}'" + (f" and name '{doc_name}'" if doc_name else "")

        result = {
            "doc_name": row["doc_name"],
            "content": row["content"],
            "parsed_data": row["parsed_data"],
            "created_at": str(row["created_at"]),
            "updated_at": str(row["updated_at"])
        }
        return json.dumps(result, indent=2)
    except Exception as e:
        return f"Error retrieving document: {str(e)}"


# =============================================================================
# TEMPLATE PARSING TOOLS (Schema Compliance Agent)
# =============================================================================

@function_tool
def parse_field_specification(field_name: str, data_type: str, nullable: bool = True,
                              default_value: Optional[str] = None,
                              agreement_types: Optional[str] = None,
                              conditional_rules: Optional[str] = None) -> str:
    """
    Parse and store a field specification from the vendor template.

    Args:
        field_name: Name of the field (e.g., 'supplier_performance_reporting')
        data_type: Data type (e.g., 'VARCHAR', 'BOOLEAN', 'INTEGER', 'DATE')
        nullable: Whether the field can be null
        default_value: Default value if applicable
        agreement_types: Comma-separated list of agreement types this field applies to
        conditional_rules: JSON string of conditional rules (e.g., '{"if_type": "amendment", "then": "required"}')

    Returns:
        Success message with parsed specification or error message.
    """
    try:
        init_database()
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO field_specifications
            (field_name, data_type, nullable, default_value, agreement_types, conditional_rules)
            VALUES (%s, %s, %s, %s, %s, %s)
            ON CONFLICT (field_name)
            DO UPDATE SET data_type = EXCLUDED.data_type, nullable = EXCLUDED.nullable,
                          default_value = EXCLUDED.default_value, agreement_types = EXCLUDED.agreement_types,
                          conditional_rules = EXCLUDED.conditional_rules
        """, (field_name, data_type.upper(), nullable, default_value, agreement_types, conditional_rules))

        conn.commit()
        _close_cursor(cursor, conn)

        spec_summary = f"Field: {field_name}, Type: {data_type.upper()}, Nullable: {nullable}"
        if default_value:
            spec_summary += f", Default: {default_value}"
        if agreement_types:
            spec_summary += f", Agreement Types: {agreement_types}"

        return f"Field specification parsed and stored. {spec_summary}"
    except Exception as e:
        return f"Error parsing field specification: {str(e)}"


@function_tool
def validate_template_compliance(field_name: str) -> str:
    """
    Validate that a field meets the vendor template requirements.
    Checks existing dbt models and mapping documents for compliance.

    Args:
        field_name: Name of the field to validate

    Returns:
        Compliance report with any mismatches found.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        # Get field specification
        cursor.execute("SELECT * FROM field_specifications WHERE field_name = %s", (field_name,))
        spec = cursor.fetchone()

        if not spec:
            _close_cursor(cursor, conn)
            return f"No specification found for field '{field_name}'. Please parse the field first."

        issues = []
        checks_passed = []

        # Check if field exists in dbt models
        cursor.execute("SELECT model_name, layer, original_sql FROM dbt_modifications")
        models = cursor.fetchall()

        for model in models:
            if model["original_sql"] and field_name.lower() in model["original_sql"].lower():
                checks_passed.append(f"Field found in {model['model_name']} ({model['layer']} layer)")
            else:
                issues.append(f"Field NOT found in {model['model_name']} ({model['layer']} layer)")

        # Check mapping document
        cursor.execute("SELECT content FROM documents WHERE doc_type = 'mapping_document' ORDER BY updated_at DESC LIMIT 1")
        mapping = cursor.fetchone()

        if mapping and mapping["content"]:
            if field_name.lower() in mapping["content"].lower():
                checks_passed.append("Field found in mapping document")
            else:
                issues.append("Field NOT found in mapping document")

        _close_cursor(cursor, conn)

        report = f"Template Compliance Report for '{field_name}':\n"
        report += f"Specification: Type={spec['data_type']}, Nullable={spec['nullable']}\n"
        report += f"\nChecks Passed ({len(checks_passed)}):\n"
        for check in checks_passed:
            report += f"  ✓ {check}\n"
        report += f"\nIssues Found ({len(issues)}):\n"
        for issue in issues:
            report += f"  ✗ {issue}\n"

        return report
    except Exception as e:
        return f"Error validating template compliance: {str(e)}"


@function_tool
def get_agreement_type_rules(agreement_type: str) -> str:
    """
    Get all field rules that apply to a specific agreement type.

    Args:
        agreement_type: The agreement type to get rules for (e.g., 'amendment', 'renewal', 'new')

    Returns:
        List of fields and their rules for the specified agreement type.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT field_name, data_type, nullable, default_value, conditional_rules
            FROM field_specifications
            WHERE agreement_types LIKE %s OR agreement_types IS NULL
        """, (f"%{agreement_type}%",))

        rows = cursor.fetchall()
        _close_cursor(cursor, conn)

        if not rows:
            return f"No field rules found for agreement type '{agreement_type}'"

        result = f"Field rules for agreement type '{agreement_type}':\n\n"
        for row in rows:
            result += f"- {row['field_name']}: {row['data_type']}"
            if not row['nullable']:
                result += " (REQUIRED)"
            if row['default_value']:
                result += f" [default: {row['default_value']}]"
            if row['conditional_rules']:
                result += f"\n  Conditions: {row['conditional_rules']}"
            result += "\n"

        return result
    except Exception as e:
        return f"Error getting agreement type rules: {str(e)}"


# =============================================================================
# MAPPING EXTRACTION TOOLS (Mapping Extraction Agent)
# =============================================================================

@function_tool
def extract_source_mapping(field_name: str, source_db: str, source_schema: str,
                           source_table: str, source_columns: str,
                           transform_sql: Optional[str] = None,
                           join_keys: Optional[str] = None,
                           notes: Optional[str] = None) -> str:
    """
    Extract and store a source-to-target mapping entry for a field.

    Args:
        field_name: Target field name
        source_db: Source database name
        source_schema: Source schema name
        source_table: Source table name
        source_columns: Comma-separated list of source column names
        transform_sql: SQL transformation logic (e.g., 'COALESCE(col1, col2)')
        join_keys: Comma-separated list of join key columns
        notes: Additional notes about the mapping

    Returns:
        Success message with mapping summary or error message.
    """
    try:
        init_database()
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO mapping_entries
            (field_name, source_db, source_schema, source_table, source_columns, transform_sql, join_keys, notes)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        """, (field_name, source_db, source_schema, source_table, source_columns, transform_sql, join_keys, notes))

        entry_id = cursor.fetchone()["id"]
        conn.commit()
        _close_cursor(cursor, conn)

        mapping_summary = f"{source_db}.{source_schema}.{source_table}.{source_columns} -> {field_name}"
        if transform_sql:
            mapping_summary += f" [Transform: {transform_sql[:50]}...]" if len(transform_sql) > 50 else f" [Transform: {transform_sql}]"

        return f"Mapping entry #{entry_id} created. {mapping_summary}"
    except Exception as e:
        return f"Error extracting source mapping: {str(e)}"


@function_tool
def get_transformation_rules(field_name: Optional[str] = None) -> str:
    """
    Get transformation rules from mapping entries.

    Args:
        field_name: Optional field name to filter by. If not provided, returns all transformations.

    Returns:
        List of transformation rules with source and target details.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        if field_name:
            cursor.execute("""
                SELECT * FROM mapping_entries WHERE field_name = %s
            """, (field_name,))
        else:
            cursor.execute("SELECT * FROM mapping_entries ORDER BY field_name")

        rows = cursor.fetchall()
        _close_cursor(cursor, conn)

        if not rows:
            return f"No transformation rules found" + (f" for field '{field_name}'" if field_name else "")

        result = "Transformation Rules:\n\n"
        for row in rows:
            result += f"Field: {row['field_name']}\n"
            result += f"  Source: {row['source_db']}.{row['source_schema']}.{row['source_table']}\n"
            result += f"  Columns: {row['source_columns']}\n"
            if row['transform_sql']:
                result += f"  Transform SQL: {row['transform_sql']}\n"
            if row['join_keys']:
                result += f"  Join Keys: {row['join_keys']}\n"
            if row['notes']:
                result += f"  Notes: {row['notes']}\n"
            result += "\n"

        return result
    except Exception as e:
        return f"Error getting transformation rules: {str(e)}"


@function_tool
def get_join_keys(source_table: str) -> str:
    """
    Get join keys used for a specific source table.

    Args:
        source_table: Name of the source table

    Returns:
        List of join keys and their associated field mappings.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT field_name, join_keys, source_columns
            FROM mapping_entries
            WHERE source_table = %s AND join_keys IS NOT NULL
        """, (source_table,))

        rows = cursor.fetchall()
        _close_cursor(cursor, conn)

        if not rows:
            return f"No join keys found for source table '{source_table}'"

        result = f"Join keys for source table '{source_table}':\n\n"
        for row in rows:
            result += f"- {row['field_name']}: {row['join_keys']} (columns: {row['source_columns']})\n"

        return result
    except Exception as e:
        return f"Error getting join keys: {str(e)}"


# =============================================================================
# DBT ANALYSIS TOOLS (dbt Code Planner Agent)
# =============================================================================

@function_tool
def analyze_dbt_model_structure(model_name: str, layer: str) -> str:
    """
    Analyze the structure of a loaded dbt model to understand its patterns.

    Args:
        model_name: Name of the dbt model to analyze
        layer: Model layer - 'prep' or 'final'

    Returns:
        Structural analysis including CTEs, field list, and patterns found.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT original_sql FROM dbt_modifications
            WHERE model_name = %s AND layer = %s
        """, (model_name, layer))

        row = cursor.fetchone()
        _close_cursor(cursor, conn)

        if not row or not row["original_sql"]:
            return f"No dbt model found for '{model_name}' ({layer} layer). Please load it first."

        sql = row["original_sql"]

        # Analyze structure
        analysis = f"Structure Analysis for {model_name} ({layer} layer):\n\n"

        # Find CTEs
        cte_pattern = r'(\w+)\s+AS\s*\('
        ctes = re.findall(cte_pattern, sql, re.IGNORECASE)
        if ctes:
            analysis += f"CTEs found ({len(ctes)}): {', '.join(ctes)}\n\n"

        # Find SELECT fields (simplified parsing)
        select_pattern = r'SELECT\s+([\s\S]*?)(?:FROM|$)'
        select_matches = re.findall(select_pattern, sql, re.IGNORECASE)
        if select_matches:
            # Get fields from the main SELECT
            main_select = select_matches[-1] if select_matches else ""
            fields = [f.strip().split()[-1].strip(',') for f in main_select.split(',') if f.strip()]
            analysis += f"Fields in final SELECT ({len(fields)}): {', '.join(fields[:10])}"
            if len(fields) > 10:
                analysis += f"... and {len(fields) - 10} more"
            analysis += "\n\n"

        # Find JOINs
        join_pattern = r'(LEFT|RIGHT|INNER|FULL|CROSS)?\s*JOIN\s+(\w+)'
        joins = re.findall(join_pattern, sql, re.IGNORECASE)
        if joins:
            analysis += f"JOINs found ({len(joins)}):\n"
            for join_type, table in joins:
                analysis += f"  - {join_type or 'INNER'} JOIN {table}\n"
            analysis += "\n"

        # Check for Jinja/dbt macros
        jinja_pattern = r'\{\{.*?\}\}'
        jinja_matches = re.findall(jinja_pattern, sql)
        if jinja_matches:
            analysis += f"dbt/Jinja macros found ({len(jinja_matches)}): {', '.join(set(jinja_matches[:5]))}\n"

        # Line count
        analysis += f"\nTotal lines: {len(sql.splitlines())}"

        return analysis
    except Exception as e:
        return f"Error analyzing dbt model: {str(e)}"


@function_tool
def identify_field_insertion_points(model_name: str, layer: str, new_field_name: str) -> str:
    """
    Identify the best locations to insert a new field in a dbt model.

    Args:
        model_name: Name of the dbt model
        layer: Model layer - 'prep' or 'final'
        new_field_name: Name of the new field to insert

    Returns:
        Recommended insertion points with line numbers and context.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT original_sql FROM dbt_modifications
            WHERE model_name = %s AND layer = %s
        """, (model_name, layer))

        row = cursor.fetchone()
        _close_cursor(cursor, conn)

        if not row or not row["original_sql"]:
            return f"No dbt model found for '{model_name}' ({layer} layer)"

        sql = row["original_sql"]
        lines = sql.splitlines()

        result = f"Insertion points for '{new_field_name}' in {model_name} ({layer}):\n\n"

        insertion_points = []

        for i, line in enumerate(lines):
            # Look for end of SELECT field lists (before FROM)
            if re.match(r'^\s*FROM\s+', line, re.IGNORECASE):
                insertion_points.append({
                    "line": i,
                    "type": "before_from",
                    "context": lines[max(0, i-2):i+1]
                })

            # Look for similar fields (alphabetical insertion)
            field_match = re.match(r'^\s*,?\s*(\w+)\s*$', line)
            if field_match:
                existing_field = field_match.group(1).lower()
                if existing_field > new_field_name.lower():
                    insertion_points.append({
                        "line": i,
                        "type": "alphabetical",
                        "context": lines[max(0, i-1):i+2]
                    })
                    break

        if insertion_points:
            for point in insertion_points[:3]:
                result += f"Option ({point['type']}) - Line {point['line'] + 1}:\n"
                for ctx_line in point['context']:
                    result += f"  {ctx_line}\n"
                result += "\n"
        else:
            result += "No clear insertion points found. Manual review recommended.\n"
            result += f"Model has {len(lines)} lines total."

        return result
    except Exception as e:
        return f"Error identifying insertion points: {str(e)}"


@function_tool
def get_existing_field_patterns(model_name: str, layer: str) -> str:
    """
    Extract patterns used for existing fields in the dbt model.
    Helps ensure new fields follow the same conventions.

    Args:
        model_name: Name of the dbt model
        layer: Model layer - 'prep' or 'final'

    Returns:
        Common patterns for field definitions, aliases, and transformations.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT original_sql FROM dbt_modifications
            WHERE model_name = %s AND layer = %s
        """, (model_name, layer))

        row = cursor.fetchone()
        _close_cursor(cursor, conn)

        if not row or not row["original_sql"]:
            return f"No dbt model found for '{model_name}' ({layer} layer)"

        sql = row["original_sql"]

        result = f"Field patterns in {model_name} ({layer}):\n\n"

        # Pattern: simple field reference
        simple_fields = re.findall(r'^\s*,?\s*(\w+)\s*$', sql, re.MULTILINE)
        if simple_fields:
            result += f"Simple fields: {len(simple_fields)} found\n"
            result += f"  Examples: {', '.join(simple_fields[:5])}\n\n"

        # Pattern: aliased fields
        alias_pattern = r'(\w+(?:\.\w+)?)\s+AS\s+(\w+)'
        aliases = re.findall(alias_pattern, sql, re.IGNORECASE)
        if aliases:
            result += f"Aliased fields: {len(aliases)} found\n"
            for src, alias in aliases[:5]:
                result += f"  {src} AS {alias}\n"
            result += "\n"

        # Pattern: CASE expressions
        case_count = len(re.findall(r'\bCASE\b', sql, re.IGNORECASE))
        if case_count:
            result += f"CASE expressions: {case_count} found\n\n"

        # Pattern: COALESCE
        coalesce_count = len(re.findall(r'\bCOALESCE\b', sql, re.IGNORECASE))
        if coalesce_count:
            result += f"COALESCE expressions: {coalesce_count} found\n\n"

        # Pattern: NULL handling
        nullif_count = len(re.findall(r'\bNULLIF\b', sql, re.IGNORECASE))
        ifnull_count = len(re.findall(r'\bIFNULL\b', sql, re.IGNORECASE))
        if nullif_count or ifnull_count:
            result += f"NULL handling: NULLIF={nullif_count}, IFNULL={ifnull_count}\n\n"

        # Indentation pattern
        indent_matches = re.findall(r'^(\s+),?\s*\w+', sql, re.MULTILINE)
        if indent_matches:
            common_indent = max(set(indent_matches), key=indent_matches.count)
            result += f"Common indentation: {len(common_indent)} spaces\n"

        return result
    except Exception as e:
        return f"Error getting field patterns: {str(e)}"


# =============================================================================
# SQL GENERATION TOOLS (dbt SQL Patch Generator Agent)
# =============================================================================

@function_tool
def generate_prep_layer_sql(field_name: str, source_expression: str,
                            model_name: str, position_hint: Optional[str] = None) -> str:
    """
    Generate SQL for adding a field to the prep layer dbt model.

    Args:
        field_name: Name of the new field
        source_expression: SQL expression for the field (e.g., 'src.supplier_perf_flag')
        model_name: Name of the dbt model
        position_hint: Optional hint for where to insert ('alphabetical', 'end', or after a specific field)

    Returns:
        Generated SQL snippet and suggested insertion location.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT original_sql FROM dbt_modifications
            WHERE model_name = %s AND layer = 'prep'
        """, (model_name,))

        row = cursor.fetchone()

        if not row or not row["original_sql"]:
            _close_cursor(cursor, conn)
            return f"No prep layer model found for '{model_name}'"

        original_sql = row["original_sql"]

        # Generate the field line
        field_line = f"    , {source_expression} AS {field_name}"

        result = f"Generated SQL for prep layer ({model_name}):\n\n"
        result += f"```sql\n{field_line}\n```\n\n"

        # Find insertion location
        lines = original_sql.splitlines()
        insert_line = None

        for i, line in enumerate(lines):
            if re.match(r'^\s*FROM\s+', line, re.IGNORECASE):
                insert_line = i
                break

        if insert_line:
            result += f"Suggested insertion: Before line {insert_line + 1} (before FROM clause)\n"
            result += f"Context:\n"
            for j in range(max(0, insert_line - 3), insert_line + 1):
                result += f"  {j + 1}: {lines[j]}\n"

        # Store the modification plan (position_hint is used for documentation purposes)
        cursor.execute("""
            UPDATE dbt_modifications
            SET modified_sql = %s, status = 'sql_generated'
            WHERE model_name = %s AND layer = 'prep'
        """, (f"-- ADD: {field_line}", model_name))

        conn.commit()
        _close_cursor(cursor, conn)

        return result
    except Exception as e:
        return f"Error generating prep layer SQL: {str(e)}"


@function_tool
def generate_final_layer_sql(field_name: str, source_expression: str,
                             model_name: str, data_type: Optional[str] = None) -> str:
    """
    Generate SQL for adding a field to the final/staging layer dbt model.

    Args:
        field_name: Name of the new field
        source_expression: SQL expression for the field (e.g., 'prep.supplier_performance_reporting')
        model_name: Name of the dbt model
        data_type: Optional data type for casting (e.g., 'VARCHAR', 'BOOLEAN')

    Returns:
        Generated SQL snippet and suggested insertion location.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT original_sql FROM dbt_modifications
            WHERE model_name = %s AND layer = 'final'
        """, (model_name,))

        row = cursor.fetchone()

        if not row or not row["original_sql"]:
            _close_cursor(cursor, conn)
            return f"No final layer model found for '{model_name}'"

        original_sql = row["original_sql"]

        # Generate the field line with optional casting
        if data_type:
            field_line = f"    , CAST({source_expression} AS {data_type}) AS {field_name}"
        else:
            field_line = f"    , {source_expression} AS {field_name}"

        result = f"Generated SQL for final layer ({model_name}):\n\n"
        result += f"```sql\n{field_line}\n```\n\n"

        # Find insertion location
        lines = original_sql.splitlines()
        insert_line = None

        for i, line in enumerate(lines):
            if re.match(r'^\s*FROM\s+', line, re.IGNORECASE):
                insert_line = i
                break

        if insert_line:
            result += f"Suggested insertion: Before line {insert_line + 1} (before FROM clause)\n"
            result += f"Context:\n"
            for j in range(max(0, insert_line - 3), insert_line + 1):
                result += f"  {j + 1}: {lines[j]}\n"

        # Store the modification plan
        cursor.execute("""
            UPDATE dbt_modifications
            SET modified_sql = %s, status = 'sql_generated'
            WHERE model_name = %s AND layer = 'final'
        """, (f"-- ADD: {field_line}", model_name))

        conn.commit()
        _close_cursor(cursor, conn)

        return result
    except Exception as e:
        return f"Error generating final layer SQL: {str(e)}"


@function_tool
def create_sql_diff(model_name: str, layer: str, new_sql: str) -> str:
    """
    Create a diff between original and modified SQL for a dbt model.

    Args:
        model_name: Name of the dbt model
        layer: Model layer - 'prep' or 'final'
        new_sql: The complete modified SQL

    Returns:
        Unified diff showing changes.
    """
    try:
        import difflib

        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT original_sql FROM dbt_modifications
            WHERE model_name = %s AND layer = %s
        """, (model_name, layer))

        row = cursor.fetchone()

        if not row or not row["original_sql"]:
            _close_cursor(cursor, conn)
            return f"No original SQL found for '{model_name}' ({layer} layer)"

        original = row["original_sql"].splitlines(keepends=True)
        modified = new_sql.splitlines(keepends=True)

        diff = difflib.unified_diff(
            original, modified,
            fromfile=f"{model_name}_{layer}_original.sql",
            tofile=f"{model_name}_{layer}_modified.sql"
        )

        diff_content = ''.join(diff)

        # Store the diff
        cursor.execute("""
            UPDATE dbt_modifications
            SET modified_sql = %s, diff_content = %s, status = 'diff_created'
            WHERE model_name = %s AND layer = %s
        """, (new_sql, diff_content, model_name, layer))

        conn.commit()
        _close_cursor(cursor, conn)

        if not diff_content:
            return "No differences found between original and modified SQL."

        return f"SQL Diff for {model_name} ({layer} layer):\n\n```diff\n{diff_content}\n```"
    except Exception as e:
        return f"Error creating SQL diff: {str(e)}"


# =============================================================================
# VALIDATION TOOLS (Validator/Reviewer Agent)
# =============================================================================

@function_tool
def validate_field_in_models(field_name: str) -> str:
    """
    Validate that a field exists in both prep and final layer models.

    Args:
        field_name: Name of the field to validate

    Returns:
        Validation result with details about where the field was found or missing.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("SELECT model_name, layer, original_sql, modified_sql FROM dbt_modifications")
        models = cursor.fetchall()

        results = []
        for model in models:
            sql_to_check = model["modified_sql"] or model["original_sql"]
            if sql_to_check:
                found = field_name.lower() in sql_to_check.lower()
                results.append({
                    "model": model["model_name"],
                    "layer": model["layer"],
                    "found": found
                })

        # Store validation result
        status = "pass" if all(r["found"] for r in results) else "fail"
        cursor.execute("""
            INSERT INTO validation_results (field_name, check_type, status, message)
            VALUES (%s, 'field_existence', %s, %s)
        """, (field_name, status, json.dumps(results)))

        conn.commit()
        _close_cursor(cursor, conn)

        report = f"Field Existence Validation for '{field_name}':\n\n"
        for r in results:
            icon = "✓" if r["found"] else "✗"
            report += f"{icon} {r['model']} ({r['layer']}): {'Found' if r['found'] else 'NOT FOUND'}\n"

        report += f"\nOverall: {'PASS' if status == 'pass' else 'FAIL'}"
        return report
    except Exception as e:
        return f"Error validating field in models: {str(e)}"


@function_tool
def check_template_requirements(field_name: str) -> str:
    """
    Check that a field meets all vendor template requirements.

    Args:
        field_name: Name of the field to check

    Returns:
        Detailed requirements checklist with pass/fail status.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        # Get field specification
        cursor.execute("SELECT * FROM field_specifications WHERE field_name = %s", (field_name,))
        spec = cursor.fetchone()

        if not spec:
            _close_cursor(cursor, conn)
            return f"No specification found for field '{field_name}'"

        checks = []

        # Check 1: Field has defined data type
        checks.append({
            "check": "Data type defined",
            "pass": bool(spec["data_type"]),
            "detail": spec["data_type"] or "Missing"
        })

        # Check 2: Nullability is specified
        checks.append({
            "check": "Nullability specified",
            "pass": spec["nullable"] is not None,
            "detail": f"Nullable={spec['nullable']}"
        })

        # Check 3: Check if mapping exists
        cursor.execute("SELECT COUNT(*) as cnt FROM mapping_entries WHERE field_name = %s", (field_name,))
        mapping_count = cursor.fetchone()["cnt"]
        checks.append({
            "check": "Source mapping defined",
            "pass": mapping_count > 0,
            "detail": f"{mapping_count} mapping(s) found"
        })

        # Check 4: Check if field in dbt models
        cursor.execute("""
            SELECT COUNT(*) as cnt FROM dbt_modifications
            WHERE (original_sql LIKE %s OR modified_sql LIKE %s)
        """, (f"%{field_name}%", f"%{field_name}%"))
        dbt_count = cursor.fetchone()["cnt"]
        checks.append({
            "check": "Field in dbt models",
            "pass": dbt_count > 0,
            "detail": f"Found in {dbt_count} model(s)"
        })

        # Store result
        all_pass = all(c["pass"] for c in checks)
        cursor.execute("""
            INSERT INTO validation_results (field_name, check_type, status, message)
            VALUES (%s, 'template_requirements', %s, %s)
        """, (field_name, "pass" if all_pass else "fail", json.dumps(checks)))

        conn.commit()
        _close_cursor(cursor, conn)

        report = f"Template Requirements Check for '{field_name}':\n\n"
        for c in checks:
            icon = "✓" if c["pass"] else "✗"
            report += f"{icon} {c['check']}: {c['detail']}\n"

        report += f"\nOverall: {'PASS' if all_pass else 'FAIL'}"
        return report
    except Exception as e:
        return f"Error checking template requirements: {str(e)}"


@function_tool
def verify_contract_constraint(model_name: str) -> str:
    """
    Verify that the one-row-per-contract constraint is maintained.
    Checks for proper GROUP BY, DISTINCT, or unique key handling.

    Args:
        model_name: Name of the dbt model to check

    Returns:
        Constraint verification result.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT layer, original_sql, modified_sql FROM dbt_modifications
            WHERE model_name = %s
        """, (model_name,))

        rows = cursor.fetchall()
        _close_cursor(cursor, conn)

        if not rows:
            return f"No model found for '{model_name}'"

        report = f"Contract Constraint Verification for '{model_name}':\n\n"

        for row in rows:
            sql = row["modified_sql"] or row["original_sql"]
            if not sql:
                continue

            layer = row["layer"]
            findings = []

            # Check for GROUP BY
            if re.search(r'\bGROUP\s+BY\b', sql, re.IGNORECASE):
                findings.append("GROUP BY clause found")

            # Check for DISTINCT
            if re.search(r'\bDISTINCT\b', sql, re.IGNORECASE):
                findings.append("DISTINCT keyword found")

            # Check for QUALIFY (Snowflake)
            if re.search(r'\bQUALIFY\b', sql, re.IGNORECASE):
                findings.append("QUALIFY clause found (deduplication)")

            # Check for ROW_NUMBER
            if re.search(r'\bROW_NUMBER\b', sql, re.IGNORECASE):
                findings.append("ROW_NUMBER function found")

            # Check for contract_id in SELECT
            if re.search(r'\bcontract_id\b', sql, re.IGNORECASE):
                findings.append("contract_id field referenced")

            report += f"{layer.upper()} layer:\n"
            if findings:
                for f in findings:
                    report += f"  ✓ {f}\n"
            else:
                report += "  ⚠ No explicit uniqueness handling found\n"
            report += "\n"

        return report
    except Exception as e:
        return f"Error verifying contract constraint: {str(e)}"


@function_tool
def generate_validation_checklist(field_name: str) -> str:
    """
    Generate a comprehensive validation checklist for a field migration.

    Args:
        field_name: Name of the field being migrated

    Returns:
        Complete checklist with all validation steps and their status.
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        checklist = []

        # 1. Field specification exists
        cursor.execute("SELECT COUNT(*) as cnt FROM field_specifications WHERE field_name = %s", (field_name,))
        spec_exists = cursor.fetchone()["cnt"] > 0
        checklist.append(("Field specification defined", spec_exists))

        # 2. Source mapping exists
        cursor.execute("SELECT COUNT(*) as cnt FROM mapping_entries WHERE field_name = %s", (field_name,))
        mapping_exists = cursor.fetchone()["cnt"] > 0
        checklist.append(("Source mapping documented", mapping_exists))

        # 3. Prep layer updated
        cursor.execute("""
            SELECT modified_sql FROM dbt_modifications
            WHERE layer = 'prep' AND (modified_sql LIKE %s OR original_sql LIKE %s)
        """, (f"%{field_name}%", f"%{field_name}%"))
        prep_updated = cursor.fetchone() is not None
        checklist.append(("Prep layer model updated", prep_updated))

        # 4. Final layer updated
        cursor.execute("""
            SELECT modified_sql FROM dbt_modifications
            WHERE layer = 'final' AND (modified_sql LIKE %s OR original_sql LIKE %s)
        """, (f"%{field_name}%", f"%{field_name}%"))
        final_updated = cursor.fetchone() is not None
        checklist.append(("Final layer model updated", final_updated))

        # 5. Validation results exist
        cursor.execute("SELECT COUNT(*) as cnt FROM validation_results WHERE field_name = %s", (field_name,))
        validated = cursor.fetchone()["cnt"] > 0
        checklist.append(("Validation checks completed", validated))

        _close_cursor(cursor, conn)

        report = f"Validation Checklist for '{field_name}':\n\n"
        all_pass = True
        for item, passed in checklist:
            icon = "✓" if passed else "☐"
            report += f"{icon} {item}\n"
            if not passed:
                all_pass = False

        report += f"\n{'='*40}\n"
        report += f"Status: {'COMPLETE' if all_pass else 'INCOMPLETE'}\n"
        report += f"Items passed: {sum(1 for _, p in checklist if p)}/{len(checklist)}"

        return report
    except Exception as e:
        return f"Error generating validation checklist: {str(e)}"


# =============================================================================
# TESTING TOOLS (Regression Diff + Test Agent)
# =============================================================================

@function_tool
def generate_dbt_tests(field_name: str, data_type: str, nullable: bool = True,
                       accepted_values: Optional[str] = None) -> str:
    """
    Generate dbt test definitions for a new field.

    Args:
        field_name: Name of the field to test
        data_type: Data type of the field
        nullable: Whether the field can be null
        accepted_values: Comma-separated list of accepted values (for enum-like fields)

    Returns:
        Generated dbt test YAML configuration.
    """
    try:
        tests = []

        # Not null test
        if not nullable:
            tests.append(f"      - not_null")

        # Accepted values test
        if accepted_values:
            values = [v.strip() for v in accepted_values.split(',')]
            values_yaml = "\n".join([f"            - '{v}'" for v in values])
            tests.append(f"      - accepted_values:\n          values:\n{values_yaml}")

        # Type-specific tests
        if data_type.upper() in ('DATE', 'TIMESTAMP'):
            tests.append(f"      # Consider adding: dbt_utils.expression_is_true for date validation")
        elif data_type.upper() in ('INTEGER', 'NUMBER', 'DECIMAL'):
            tests.append(f"      # Consider adding: dbt_utils.expression_is_true for range validation")

        yaml_content = f"""  - name: {field_name}
    description: "TODO: Add description"
    tests:
{chr(10).join(tests) if tests else "      # No tests defined"}
"""

        result = f"Generated dbt tests for '{field_name}':\n\n```yaml\n{yaml_content}\n```\n\n"
        result += "Add this to your schema.yml file under the appropriate model's columns section."

        return result
    except Exception as e:
        return f"Error generating dbt tests: {str(e)}"


@function_tool
def create_validation_query(field_name: str, source_table: str, target_table: str,
                            join_key: str = "contract_id") -> str:
    """
    Create a validation query to compare source and target field values.

    Args:
        field_name: Name of the field to validate
        source_table: Fully qualified source table name
        target_table: Fully qualified target table name
        join_key: Column to join source and target (default: contract_id)

    Returns:
        SQL validation query for data comparison.
    """
    try:
        query = f"""-- Validation Query for {field_name}
-- Purpose: Compare source and target values to ensure migration accuracy

-- Count comparison
SELECT
    'source' as dataset,
    COUNT(*) as total_rows,
    COUNT({field_name}) as non_null_count,
    COUNT(*) - COUNT({field_name}) as null_count
FROM {source_table}

UNION ALL

SELECT
    'target' as dataset,
    COUNT(*) as total_rows,
    COUNT({field_name}) as non_null_count,
    COUNT(*) - COUNT({field_name}) as null_count
FROM {target_table};

-- Value mismatch check
SELECT
    s.{join_key},
    s.{field_name} as source_value,
    t.{field_name} as target_value,
    CASE
        WHEN s.{field_name} = t.{field_name} THEN 'MATCH'
        WHEN s.{field_name} IS NULL AND t.{field_name} IS NULL THEN 'BOTH_NULL'
        ELSE 'MISMATCH'
    END as comparison_result
FROM {source_table} s
LEFT JOIN {target_table} t ON s.{join_key} = t.{join_key}
WHERE s.{field_name} != t.{field_name}
   OR (s.{field_name} IS NULL AND t.{field_name} IS NOT NULL)
   OR (s.{field_name} IS NOT NULL AND t.{field_name} IS NULL)
LIMIT 100;

-- Distinct value comparison
SELECT 'source' as dataset, {field_name}, COUNT(*) as cnt
FROM {source_table}
GROUP BY {field_name}

UNION ALL

SELECT 'target' as dataset, {field_name}, COUNT(*) as cnt
FROM {target_table}
GROUP BY {field_name}
ORDER BY dataset, {field_name};
"""

        return f"Validation Query:\n\n```sql\n{query}\n```"
    except Exception as e:
        return f"Error creating validation query: {str(e)}"


# =============================================================================
# OUTPUT FILE TOOLS (Save and list output files)
# =============================================================================

# Output folder path
OUTPUT_DIR = Path(__file__).parent.parent / "output"


@function_tool
def save_sql_output(model_name: str, layer: str, sql_content: str, description: str = "") -> str:
    """
    Save generated SQL to the output folder.
    The file will be saved in output/sql/ with a timestamp.

    Args:
        model_name: Name of the dbt model
        layer: Model layer (prep/final)
        sql_content: SQL content to save
        description: Optional description of what this SQL does

    Returns:
        Success message with file path.
    """
    _log("file_op", f"Saving SQL: {model_name} ({layer})", "Writing to output/sql/")
    try:
        from datetime import datetime

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{model_name}_{layer}_{timestamp}.sql"

        sql_folder = OUTPUT_DIR / "sql"
        sql_folder.mkdir(parents=True, exist_ok=True)
        filepath = sql_folder / filename

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"-- Generated: {datetime.now().isoformat()}\n")
            f.write(f"-- Model: {model_name}\n")
            f.write(f"-- Layer: {layer}\n")
            if description:
                f.write(f"-- Description: {description}\n")
            f.write("-- " + "=" * 60 + "\n\n")
            f.write(sql_content)

        _log("success", f"SQL saved: {filename}")
        return f"SQL saved successfully!\n\nFile: {filename}\nLocation: output/sql/{filename}\n\nYou can find this file in the output/sql/ folder."
    except Exception as e:
        _log("error", f"Failed to save SQL: {str(e)}")
        return f"Error saving SQL file: {str(e)}"


@function_tool
def save_diff_output(model_name: str, layer: str, diff_content: str) -> str:
    """
    Save a SQL diff to the output folder.
    The file will be saved in output/diffs/ with a timestamp.

    Args:
        model_name: Name of the dbt model
        layer: Model layer (prep/final)
        diff_content: Unified diff content

    Returns:
        Success message with file path.
    """
    _log("file_op", f"Saving diff: {model_name} ({layer})", "Writing to output/diffs/")
    try:
        from datetime import datetime

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{model_name}_{layer}_diff_{timestamp}.diff"

        diff_folder = OUTPUT_DIR / "diffs"
        diff_folder.mkdir(parents=True, exist_ok=True)
        filepath = diff_folder / filename

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"# Diff for {model_name} ({layer} layer)\n")
            f.write(f"# Generated: {datetime.now().isoformat()}\n\n")
            f.write(diff_content)

        _log("success", f"Diff saved: {filename}")
        return f"Diff saved successfully!\n\nFile: {filename}\nLocation: output/diffs/{filename}\n\nYou can find this file in the output/diffs/ folder."
    except Exception as e:
        _log("error", f"Failed to save diff: {str(e)}")
        return f"Error saving diff file: {str(e)}"


@function_tool
def save_validation_report_output(field_name: str, report_content: str) -> str:
    """
    Save a validation report to the output folder.
    The file will be saved in output/reports/ with a timestamp.

    Args:
        field_name: Name of the field being validated
        report_content: Validation report content

    Returns:
        Success message with file path.
    """
    try:
        from datetime import datetime

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"validation_{field_name}_{timestamp}.txt"

        reports_folder = OUTPUT_DIR / "reports"
        reports_folder.mkdir(parents=True, exist_ok=True)
        filepath = reports_folder / filename

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"VALIDATION REPORT\n")
            f.write(f"=" * 60 + "\n")
            f.write(f"Field: {field_name}\n")
            f.write(f"Generated: {datetime.now().isoformat()}\n")
            f.write(f"=" * 60 + "\n\n")
            f.write(report_content)

        return f"Validation report saved!\n\nFile: {filename}\nLocation: output/reports/{filename}\n\nYou can find this file in the output/reports/ folder."
    except Exception as e:
        return f"Error saving validation report: {str(e)}"


@function_tool
def save_dbt_test_output(field_name: str, yaml_content: str) -> str:
    """
    Save dbt test YAML to the output folder.
    The file will be saved in output/tests/ with a timestamp.

    Args:
        field_name: Name of the field being tested
        yaml_content: YAML content for dbt tests

    Returns:
        Success message with file path.
    """
    try:
        from datetime import datetime

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"test_{field_name}_{timestamp}.yml"

        tests_folder = OUTPUT_DIR / "tests"
        tests_folder.mkdir(parents=True, exist_ok=True)
        filepath = tests_folder / filename

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"# dbt tests for {field_name}\n")
            f.write(f"# Generated: {datetime.now().isoformat()}\n\n")
            f.write(yaml_content)

        return f"dbt test YAML saved!\n\nFile: {filename}\nLocation: output/tests/{filename}\n\nYou can find this file in the output/tests/ folder."
    except Exception as e:
        return f"Error saving dbt test file: {str(e)}"


@function_tool
def save_mapping_output(field_name: str, mapping_content: str) -> str:
    """
    Save a field mapping document to the output folder.
    The file will be saved in output/mappings/ with a timestamp.

    Args:
        field_name: Name of the field being mapped
        mapping_content: Mapping information content

    Returns:
        Success message with file path.
    """
    try:
        from datetime import datetime

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"mapping_{field_name}_{timestamp}.txt"

        mappings_folder = OUTPUT_DIR / "mappings"
        mappings_folder.mkdir(parents=True, exist_ok=True)
        filepath = mappings_folder / filename

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"FIELD MAPPING DOCUMENT\n")
            f.write(f"=" * 60 + "\n")
            f.write(f"Field: {field_name}\n")
            f.write(f"Generated: {datetime.now().isoformat()}\n")
            f.write(f"=" * 60 + "\n\n")
            f.write(mapping_content)

        return f"Mapping document saved!\n\nFile: {filename}\nLocation: output/mappings/{filename}\n\nYou can find this file in the output/mappings/ folder."
    except Exception as e:
        return f"Error saving mapping document: {str(e)}"


@function_tool
def save_migration_summary_output(field_name: str, summary_content: str) -> str:
    """
    Save a complete migration summary to the output folder.
    The file will be saved in output/summaries/ with a timestamp.

    Args:
        field_name: Name of the field being migrated
        summary_content: Complete migration summary

    Returns:
        Success message with file path.
    """
    try:
        from datetime import datetime

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"migration_summary_{field_name}_{timestamp}.txt"

        summaries_folder = OUTPUT_DIR / "summaries"
        summaries_folder.mkdir(parents=True, exist_ok=True)
        filepath = summaries_folder / filename

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"MIGRATION SUMMARY\n")
            f.write(f"=" * 60 + "\n")
            f.write(f"Field: {field_name}\n")
            f.write(f"Generated: {datetime.now().isoformat()}\n")
            f.write(f"=" * 60 + "\n\n")
            f.write(summary_content)

        return f"Migration summary saved!\n\nFile: {filename}\nLocation: output/summaries/{filename}\n\nYou can find this file in the output/summaries/ folder."
    except Exception as e:
        return f"Error saving migration summary: {str(e)}"


@function_tool
def list_output_files(folder: str = "") -> str:
    """
    List all output files in the output folder.

    Args:
        folder: Optional subfolder to list (sql, diffs, reports, tests, mappings, summaries, logs).
                If empty, lists all subfolders with file counts.

    Returns:
        List of output files with their details.
    """
    try:
        if folder:
            # List specific subfolder
            target_folder = OUTPUT_DIR / folder
            if not target_folder.exists():
                return f"Folder 'output/{folder}' does not exist."

            files = []
            for f in target_folder.iterdir():
                if f.is_file() and not f.name.startswith('.') and not f.name.startswith('__'):
                    stat = f.stat()
                    files.append({
                        "name": f.name,
                        "size": stat.st_size,
                        "modified": stat.st_mtime
                    })

            if not files:
                return f"No files in output/{folder}/ folder."

            # Sort by modification time (newest first)
            files.sort(key=lambda x: x["modified"], reverse=True)

            from datetime import datetime
            result = f"Output files in output/{folder}/ ({len(files)} files):\n\n"
            for f in files[:20]:  # Show max 20 files
                size_kb = f["size"] / 1024
                mod_time = datetime.fromtimestamp(f["modified"]).strftime("%Y-%m-%d %H:%M")
                result += f"- {f['name']} ({size_kb:.1f} KB, {mod_time})\n"

            if len(files) > 20:
                result += f"\n... and {len(files) - 20} more files"

            result += f"\n\nFolder path: output/{folder}/"
            return result
        else:
            # List all subfolders with counts
            subfolders = ['sql', 'diffs', 'reports', 'tests', 'mappings', 'summaries', 'logs']
            result = "Output folder structure:\n\n"

            for subfolder in subfolders:
                folder_path = OUTPUT_DIR / subfolder
                if folder_path.exists():
                    count = len([f for f in folder_path.iterdir() if f.is_file() and not f.name.startswith('.')])
                    result += f"- output/{subfolder}/ ({count} files)\n"
                else:
                    result += f"- output/{subfolder}/ (empty)\n"

            result += "\n\nUse list_output_files('sql') to see files in a specific folder."
            result += "\n\nOutput folder location: output/"
            return result
    except Exception as e:
        return f"Error listing output files: {str(e)}"


@function_tool
def get_output_folder_path() -> str:
    """
    Get the path to the output folder where all generated files are saved.

    Returns:
        Path information for the output folder.
    """
    return f"""Output Folder Location
======================

Main output folder: output/

Subfolders:
- output/sql/       - Generated SQL files
- output/diffs/     - SQL diff files showing changes
- output/reports/   - Validation reports
- output/tests/     - dbt test YAML files
- output/mappings/  - Field mapping documents
- output/summaries/ - Migration summaries
- output/logs/      - Conversation logs

All generated files are automatically saved with timestamps.
You can manually access these files at: {OUTPUT_DIR}"""


@function_tool
def read_output_file(folder: str, filename: str) -> str:
    """
    Read and return the contents of an output file.
    Use this to view generated SQL, diffs, reports, tests, or other outputs.

    Args:
        folder: The subfolder (sql, diffs, reports, tests, mappings, summaries, logs)
        filename: Name of the file to read

    Returns:
        Contents of the file.
    """
    _log("file_op", f"Reading output file: {folder}/{filename}")
    try:
        valid_folders = ['sql', 'diffs', 'reports', 'tests', 'mappings', 'summaries', 'logs']
        if folder not in valid_folders:
            return f"Invalid folder. Choose from: {', '.join(valid_folders)}"

        file_path = OUTPUT_DIR / folder / filename

        if not file_path.exists():
            # List available files in that folder
            folder_path = OUTPUT_DIR / folder
            if folder_path.exists():
                files = [f.name for f in folder_path.iterdir() if f.is_file() and not f.name.startswith('.')]
                if files:
                    return f"File '{filename}' not found in output/{folder}/.\n\nAvailable files:\n" + "\n".join(f"- {f}" for f in sorted(files))
                else:
                    return f"No files in output/{folder}/ folder."
            return f"Folder output/{folder}/ does not exist."

        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Get file info
        stat = file_path.stat()
        from datetime import datetime
        mod_time = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
        size_kb = stat.st_size / 1024

        _log("success", f"Read file: {filename} ({size_kb:.1f} KB)")

        # Truncate if very large
        max_chars = 10000
        if len(content) > max_chars:
            return f"""File: {filename}
Location: output/{folder}/{filename}
Size: {size_kb:.1f} KB
Modified: {mod_time}

Content (first {max_chars} characters):
{'=' * 60}

{content[:max_chars]}

... [truncated, {len(content) - max_chars} more characters]"""

        return f"""File: {filename}
Location: output/{folder}/{filename}
Size: {size_kb:.1f} KB
Modified: {mod_time}

Content:
{'=' * 60}

{content}"""

    except Exception as e:
        _log("error", f"Failed to read file: {str(e)}")
        return f"Error reading file: {str(e)}"


@function_tool
def get_latest_output_file(folder: str) -> str:
    """
    Get the most recently created/modified file from an output folder.
    Useful for quickly accessing the latest generated SQL, report, etc.

    Args:
        folder: The subfolder (sql, diffs, reports, tests, mappings, summaries, logs)

    Returns:
        Contents of the most recent file in that folder.
    """
    _log("tool", f"Getting latest file from output/{folder}/")
    try:
        valid_folders = ['sql', 'diffs', 'reports', 'tests', 'mappings', 'summaries', 'logs']
        if folder not in valid_folders:
            return f"Invalid folder. Choose from: {', '.join(valid_folders)}"

        folder_path = OUTPUT_DIR / folder
        if not folder_path.exists():
            return f"Folder output/{folder}/ does not exist."

        files = []
        for f in folder_path.iterdir():
            if f.is_file() and not f.name.startswith('.'):
                files.append((f, f.stat().st_mtime))

        if not files:
            return f"No files in output/{folder}/ folder."

        # Get most recent file
        latest_file = max(files, key=lambda x: x[1])[0]

        # Read and return its contents
        with open(latest_file, 'r', encoding='utf-8') as f:
            content = f.read()

        from datetime import datetime
        stat = latest_file.stat()
        mod_time = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
        size_kb = stat.st_size / 1024

        _log("success", f"Found latest: {latest_file.name}")

        max_chars = 10000
        if len(content) > max_chars:
            content = content[:max_chars] + f"\n\n... [truncated, {len(content) - max_chars} more characters]"

        return f"""Latest file in output/{folder}/:

File: {latest_file.name}
Size: {size_kb:.1f} KB
Modified: {mod_time}

Content:
{'=' * 60}

{content}"""

    except Exception as e:
        _log("error", f"Failed to get latest file: {str(e)}")
        return f"Error: {str(e)}"
